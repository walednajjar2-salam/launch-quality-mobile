#!/usr/bin/env python3
"""Generate QUALITY OF LAUNCH PROJECTS LLC hospitality manager letter (.docx)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

NAVY = RGBColor(0x0D, 0x17, 0x28)
TEAL = RGBColor(0x05, 0xAA, 0xA1)
GOLD = RGBColor(0xC9, 0xA2, 0x27)
DARK = RGBColor(0x11, 0x18, 0x27)
CREAM = RGBColor(0xF7, 0xEE, 0xDC)

OUTPUT = Path(__file__).resolve().parent.parent / "documents" / "مدير-إدارة-الضيافة.docx"

RESPONSIBILITIES = [
    "إدارة الحجوزات والفعاليات",
    "متابعة تجهيز الحفلات والمناسبات",
    "الإشراف على فرق التشغيل والعملاء",
    "إدارة المعدات والعروض الخاصة بالحفلات والفعاليات",
    "متابعة جودة الخدمة المقدمة بالفعاليات",
    "استلام وتسليم معدات وتجهيزات الحفلات",
    "متابعة كميات وأصناف المعدات والتجهيزات",
    "التأكد من سلامة وصلاحية المعدات",
    "مطابقة الموجودات والتالف ورفع طلبات الشراء",
    "المحافظة على جاهزية تنظيم الفعالية مع سجلات المخزون",
    "متابعة نقل المعدات لأي فعالية",
    "التنسيق مع قسم النقليات لتوفير وسائل النقل اللازمة للفعاليات والفرق التشغيلية",
]


def set_rtl(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    p_pr.append(bidi)


def add_colored_paragraph(doc, text, *, size=11, bold=False, color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT, rtl=False):
    p = doc.add_paragraph()
    if rtl:
        set_rtl(p)
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    run.font.color.rgb = color
    return p


def add_footer_line(doc, text, icon=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"{text}  {icon}".strip())
    run.font.size = Pt(10)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.paragraph_format.space_after = Pt(2)


def build_document() -> Document:
    doc = Document()

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # Company name
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("QUALITY OF LAUNCH PROJECTS\nLLC")
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = "Arial"
    run.font.color.rgb = DARK

    subtitle = add_colored_paragraph(
        doc,
        "REAL ESTATE • PROPERTY MANAGEMENT • HOSPITALITY",
        size=10,
        bold=True,
        color=GOLD,
    )
    subtitle.paragraph_format.space_after = Pt(24)

    # Arabic job title
    job_title = add_colored_paragraph(
        doc,
        "مدير إدارة الضيافة وتنظيم الحفلات",
        size=18,
        bold=True,
        color=NAVY,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        rtl=True,
    )
    job_title.paragraph_format.space_after = Pt(16)

    # Responsibilities
    for item in RESPONSIBILITIES:
        p = doc.add_paragraph(style="List Bullet")
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.line_spacing = 1.8
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.size = Pt(13)
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:cs"), "Arial")
        run.font.color.rgb = DARK

    doc.add_paragraph()

    # Signature
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lines = ["Regards,", "Signature", "QUALITY OF LAUNCH PROJECTS LLC"]
    for i, line in enumerate(lines):
        run = sig.add_run(line + ("\n" if i < len(lines) - 1 else ""))
        run.font.size = Pt(11)
        run.font.name = "Arial"
        run.font.color.rgb = DARK
        if line == "Signature":
            run.bold = True

    doc.add_paragraph()

    # Contact footer block
    footer_heading = add_colored_paragraph(
        doc,
        "—" * 40,
        size=8,
        color=TEAL,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    footer_heading.paragraph_format.space_before = Pt(12)

    contacts = [
        ("info@qualitylaunch.om", "✉"),
        ("+968 92120205", "☎"),
        ("www.qualitylaunch.om", "🌐"),
        ("Nizwa, Sultanate of Oman", "⌖"),
    ]

    for text, icon in contacts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"{text}  {icon}")
        run.font.size = Pt(10)
        run.font.name = "Arial"
        run.font.color.rgb = TEAL
        p.paragraph_format.space_after = Pt(2)

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
