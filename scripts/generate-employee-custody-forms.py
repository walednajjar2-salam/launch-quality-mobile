#!/usr/bin/env python3
"""Generate official custody receipt forms (نموذج استلام عهدة) per employee."""

from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from employee_roster import EMPLOYEES, custody_files_dir, employee_files_dir, template_path
from letterhead_gold_docx import set_rtl, set_run_font

GOLD = RGBColor(0xB8, 0x86, 0x0B)
GOLD_TEXT = RGBColor(0x7A, 0x5C, 0x10)
DARK = RGBColor(0x2D, 0x22, 0x12)
MUTED = RGBColor(0x9C, 0x8A, 0x6E)

CUSTODY_ITEMS = 4

COMMITMENTS = [
    "المحافظة على العهدة وعدم إساءة استخدامها.",
    "عدم تسليمها أو إعارتها لأي طرف دون إذن رسمي.",
    "استخدامها فقط لأغراض العمل المكلف به.",
    "الإبلاغ فورًا عن أي تلف أو فقدان.",
    "إعادة العهدة عند انتهاء العمل أو عند طلب الإدارة.",
    "تحمل تكلفة الإصلاح أو الاستبدال في حال التلف الناتج عن سوء الاستخدام.",
    "إعادة العهدة بنفس الحالة التي استلمها بها (باستثناء الاستهلاك الطبيعي).",
]

EMPLOYEE_ACK = (
    "أقر أنا الموظف المذكور أعلاه بأنني استلمت العهد الموضحة أعلاه كاملة وبحالة جيدة، "
    "وأتحمل كامل المسؤولية عنها طوال فترة وجودها بحوزتي."
)


def output_dir() -> Path:
    return custody_files_dir()


def safe_filename(name: str, title: str) -> str:
    clean_name = re.sub(r'[\\/:*?"<>|]', "-", name.strip())
    clean_title = re.sub(r'[\\/:*?"<>|]', "-", title.strip())[:35]
    return f"عهدة — {clean_name} — {clean_title}.docx"


def add_heading(doc: Document, text: str, size: int = 16, space_before: int = 16) -> None:
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=GOLD)


def add_section_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True, color=GOLD_TEXT)


def add_body(doc: Document, text: str, *, bold: bool = False, space_after: int = 6) -> None:
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.7
    run = p.add_run(text)
    set_run_font(run, size=11, bold=bold, color=DARK)


def add_field_line(doc: Document, label: str, value: str = "", dots: int = 30) -> None:
    fill = value if value and not value.startswith("_") else "." * dots
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(8)
    lbl = p.add_run(f"{label}: ")
    set_run_font(lbl, size=11, bold=True, color=GOLD_TEXT)
    val = p.add_run(fill)
    set_run_font(val, size=11, color=DARK)


def add_meta_row(doc: Document, employee) -> None:
    fields = [
        ("اسم الموظف", employee.name if not employee.name.startswith("_") else ""),
        ("الوظيفة", employee.job_title),
        ("رقم الهوية / الإقامة", ""),
        ("رقم الهاتف", ""),
        ("تاريخ الاستلام", ""),
    ]
    for label, value in fields:
        add_field_line(doc, label, value, dots=42)


def build_custody_form(employee, dest: Path) -> None:
    shutil.copy2(template_path(), dest)
    doc = Document(str(dest))

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(4)

    add_heading(doc, "نموذج استلام عهدة – رسمي", size=17, space_before=8)
    add_meta_row(doc, employee)

    add_section_title(doc, "أولًا: بيانات العهدة المستلمة")
    add_body(
        doc,
        "يقر الموظف المذكور أعلاه بأنه استلم العهد التالية من الشركة، "
        "وهي بحالة جيدة وصالحة للاستخدام:",
    )
    for i in range(1, CUSTODY_ITEMS + 1):
        add_field_line(doc, str(i), dots=60)

    add_section_title(doc, "ثانيًا: التزامات الموظف")
    for i, item in enumerate(COMMITMENTS, 1):
        p = doc.add_paragraph()
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.left_indent = Cm(0.15)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.65
        marker = p.add_run(f"{i}. ")
        set_run_font(marker, size=10.5, bold=True, color=GOLD)
        body = p.add_run(item)
        set_run_font(body, size=10.5, color=DARK)

    add_section_title(doc, "ثالثًا: إقرار الموظف")
    add_body(doc, EMPLOYEE_ACK, space_after=10)
    add_field_line(doc, "اسم الموظف", employee.name if not employee.name.startswith("_") else "")
    add_field_line(doc, "التوقيع")
    add_field_line(doc, "التاريخ")

    add_section_title(doc, "رابعًا: اعتماد الإدارة")
    add_field_line(doc, "اسم المسؤول", "يعقوب فاضل الخصيبي")
    add_field_line(doc, "الوظيفة", "المدير العام")
    add_field_line(doc, "التوقيع")
    add_field_line(doc, "الختم")

    doc.save(str(dest))


def main() -> None:
    out = output_dir()
    created = 0
    for employee in EMPLOYEES:
        dest = out / safe_filename(employee.name, employee.job_title)
        build_custody_form(employee, dest)
        created += 1
        print(f"Created: {dest}")
    print(f"\nDone — {created} custody forms in:\n{out}")


if __name__ == "__main__":
    main()
