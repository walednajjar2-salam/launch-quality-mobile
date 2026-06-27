"""Gold-bordered official letterhead — matches QUALITY OF LAUNCH branded stationery."""

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from org_info import CONTACT_LINE_AR, CONTACT_LINE_EN, FOOTER_LINE, ORG

REPO_ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = REPO_ROOT / "documents"
SOURCE_IMG = DOCS_DIR / "gold-letterhead-source.png"
GOLD_HEADER_IMG = DOCS_DIR / "letterhead-gold-header.png"
GOLD_FOOTER_IMG = DOCS_DIR / "letterhead-gold-footer.png"
LOGO_GOLD_IMG = DOCS_DIR / "logo-gold.png"
LOGO_3D_IMG = DOCS_DIR / "logo-3d-gold.png"
LOGO_BADGE_IMG = DOCS_DIR / "logo-3d-badge.png"
HEADER_BAND_IMG = DOCS_DIR / "letterhead-header-band.png"
HEADER_SOURCE_RAW = DOCS_DIR / "letterhead-header-source.png"

# Logo placeholder region in header source (1024×561)
LOGO_SLOT = (714, 28, 996, 198)
CORNER_KEEP = (940, 0, 1024, 55)
HEADER_RENDER_WIDTH = 1654

GOLD_HEX = "B8860B"
GOLD_LIGHT = "D4AF37"
GOLD_DARK = "8B6914"
WHITE = "FFFFFF"
GOLD_RGB = RGBColor(0xB8, 0x86, 0x0B)
GOLD_TEXT = RGBColor(0x7A, 0x5C, 0x10)
DARK = RGBColor(0x2D, 0x22, 0x12)
MUTED = RGBColor(0x9C, 0x8A, 0x6E)


def set_rtl(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    p_pr.append(bidi)


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tc_pr.append(margins)


def remove_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    tbl_pr.append(borders)


def set_run_font(run, size=11, bold=False, color=DARK, name="Arial") -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:cs"), name)


def set_page_borders(section) -> None:
    sect_pr = section._sectPr
    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "double")
        border.set(qn("w:sz"), "18")
        border.set(qn("w:space"), "18")
        border.set(qn("w:color"), GOLD_HEX)
        pg_borders.append(border)
    sect_pr.append(pg_borders)


def _shape_arabic(text: str) -> str:
    try:
        import arabic_reshaper
        from bidi.algorithm import get_display

        return get_display(arabic_reshaper.reshape(text))
    except ImportError:
        return text


def _load_font(size: int, bold: bool = False, arabic: bool = False):
    from PIL import ImageFont

    if arabic:
        candidates = [
            Path("C:/Windows/Fonts/tahoma.ttf"),
            Path("C:/Windows/Fonts/arial.ttf"),
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        ]
    else:
        candidates = [
            Path("C:/Windows/Fonts/arialbd.ttf") if bold else Path("C:/Windows/Fonts/arial.ttf"),
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")
            if bold
            else Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    from PIL import ImageFont

    return ImageFont.load_default()


def _draw_centered_in_box(draw, text, y, left, box_width, font, fill):
    bbox = draw.textbbox((0, 0), text, font=font)
    tw = bbox[2] - bbox[0]
    draw.text((left + max(0, (box_width - tw) // 2), y), text, fill=fill, font=font)


def _draw_wrapped_en(draw, text, y, left, box_width, font, fill, line_height=22):
    words = text.split()
    lines = []
    current = []
    for word in words:
        trial = " ".join(current + [word])
        bbox = draw.textbbox((0, 0), trial, font=font)
        if bbox[2] - bbox[0] <= box_width - 20 or not current:
            current.append(word)
        else:
            lines.append(" ".join(current))
            current = [word]
    if current:
        lines.append(" ".join(current))
    for line in lines:
        _draw_centered_in_box(draw, line, y, left, box_width, font, fill)
        y += line_height
    return y


def _draw_right_in_box(draw, text, y, left, box_width, font, fill):
    bbox = draw.textbbox((0, 0), text, font=font)
    tw = bbox[2] - bbox[0]
    draw.text((left + max(0, box_width - tw - 8), y), text, fill=fill, font=font)


def _draw_left_in_box(draw, text, y, left, box_width, font, fill):
    draw.text((left + 8, y), text, fill=fill, font=font)


def _draw_wrapped_en_left(draw, text, y, left, box_width, font, fill, line_height=22):
    words = text.split()
    lines = []
    current = []
    for word in words:
        trial = " ".join(current + [word])
        bbox = draw.textbbox((0, 0), trial, font=font)
        if bbox[2] - bbox[0] <= box_width - 20 or not current:
            current.append(word)
        else:
            lines.append(" ".join(current))
            current = [word]
    if current:
        lines.append(" ".join(current))
    for line in lines:
        _draw_left_in_box(draw, line, y, left, box_width, font, fill)
        y += line_height
    return y


def _remove_dark_background(img, threshold: int = 42):
    from PIL import Image

    rgba = img.convert("RGBA")
    pixels = rgba.load()
    w, h = rgba.size
    for y in range(h):
        for x in range(w):
            r, g, b, a = pixels[x, y]
            if r < threshold and g < threshold and b < threshold:
                pixels[x, y] = (0, 0, 0, 0)
    return rgba


def _trim_transparent(img):
    from PIL import Image

    bbox = img.getbbox()
    if not bbox:
        return img
    return img.crop(bbox)


def _add_logo_shadow(logo, offset=(8, 10), blur_radius: int = 12, opacity: int = 70):
    from PIL import Image, ImageFilter

    shadow = Image.new("RGBA", logo.size, (0, 0, 0, 0))
    alpha = logo.split()[3]
    shadow_layer = Image.new("RGBA", logo.size, (90, 70, 20, opacity))
    shadow_layer.putalpha(alpha)
    shadow = Image.new("RGBA", (logo.width + offset[0] + blur_radius, logo.height + offset[1] + blur_radius), (0, 0, 0, 0))
    blurred = shadow_layer.filter(ImageFilter.GaussianBlur(blur_radius))
    shadow.paste(blurred, (offset[0], offset[1]), blurred)
    combined = Image.new("RGBA", shadow.size, (0, 0, 0, 0))
    combined.paste(shadow, (0, 0), shadow)
    combined.paste(logo, (0, 0), logo)
    return combined


def _draw_gold_arc_pair(draw, center_x: int, y: int, span: int, color: str) -> None:
    for side in (-1, 1):
        start = center_x + side * 28
        end = center_x + side * span
        draw.arc(
            [min(start, end) - 40, y - 18, max(start, end) + 40, y + 34],
            start=200 if side < 0 else 340,
            end=340 if side < 0 else 520,
            fill=color,
            width=2,
        )


def _build_logo_badge() -> None:
    from PIL import Image, ImageDraw

    if not LOGO_3D_IMG.exists():
        return

    logo = _remove_dark_background(Image.open(LOGO_3D_IMG))
    logo = _trim_transparent(logo)
    logo.thumbnail((340, 180), Image.LANCZOS)
    logo = _add_logo_shadow(logo)

    pad_x, pad_y = 48, 36
    badge_w = logo.width + pad_x * 2
    badge_h = logo.height + pad_y * 2 + 28
    badge = Image.new("RGBA", (badge_w, badge_h), (0, 0, 0, 0))
    draw = ImageDraw.Draw(badge)

    glow = Image.new("RGBA", (badge_w, badge_h), (0, 0, 0, 0))
    glow_draw = ImageDraw.Draw(glow)
    glow_draw.ellipse(
        [pad_x - 20, pad_y + logo.height // 2, badge_w - pad_x + 20, badge_h - 8],
        fill=(212, 175, 55, 35),
    )
    glow = glow.filter(__import__("PIL.ImageFilter", fromlist=["GaussianBlur"]).GaussianBlur(10))
    badge = Image.alpha_composite(badge, glow)
    draw = ImageDraw.Draw(badge)

    pedestal_y = badge_h - 22
    draw.line([(pad_x - 10, pedestal_y), (badge_w - pad_x + 10, pedestal_y)], fill=f"#{GOLD_LIGHT}", width=2)
    draw.line([(pad_x + 20, pedestal_y + 4), (badge_w - pad_x - 20, pedestal_y + 4)], fill=f"#{GOLD_HEX}", width=1)

    badge.paste(logo, (pad_x, pad_y - 8), logo)
    badge.save(LOGO_BADGE_IMG)


def uses_full_header() -> bool:
    return HEADER_BAND_IMG.exists() or HEADER_SOURCE_RAW.exists()


def _ensure_header_source() -> Path | None:
    if HEADER_SOURCE_RAW.exists():
        return HEADER_SOURCE_RAW
    assets_dir = Path.home() / ".cursor/projects/c-Projects-launch-quality-mobile-github/assets"
    if assets_dir.exists():
        matches = sorted(assets_dir.glob("*ChatGPT_Image*header*.png"))
        if not matches:
            matches = sorted(assets_dir.glob("*ChatGPT_Image*.png"))
        if matches:
            DOCS_DIR.mkdir(parents=True, exist_ok=True)
            import shutil

            shutil.copy2(matches[-1], HEADER_SOURCE_RAW)
            return HEADER_SOURCE_RAW
    return None


def _clean_header_logo_slot(img):
    """Remove 'مكان الشعار' placeholder and its frame, keep gold corner."""
    x1, y1, x2, y2 = LOGO_SLOT
    cx1, cy1, cx2, cy2 = CORNER_KEEP

    try:
        import cv2
        import numpy as np

        rgb = img.convert("RGB")
        arr = cv2.cvtColor(np.array(rgb), cv2.COLOR_RGB2BGR)
        h, w = arr.shape[:2]
        mask = np.zeros((h, w), np.uint8)
        cv2.rectangle(mask, (x1, y1), (x2, y2), 255, -1)
        cv2.rectangle(mask, (cx1, cy1), (cx2, cy2), 0, -1)
        cleaned = cv2.inpaint(arr, mask, 5, cv2.INPAINT_TELEA)
        from PIL import Image

        return Image.fromarray(cv2.cvtColor(cleaned, cv2.COLOR_BGR2RGB))
    except ImportError:
        from PIL import Image, ImageDraw, ImageFilter

        out = img.convert("RGBA").copy()
        slot_w, slot_h = x2 - x1, y2 - y1
        samples = [out.getpixel((p, p))[:3] for p in range(8, 36, 4)]
        r = sum(s[0] for s in samples) // len(samples)
        g = sum(s[1] for s in samples) // len(samples)
        b = sum(s[2] for s in samples) // len(samples)
        patch = Image.new("RGBA", (slot_w, slot_h), (r, g, b, 255))
        mask = Image.new("L", (slot_w, slot_h), 255)
        mask_draw = ImageDraw.Draw(mask)
        mask_draw.rectangle([0, 0, slot_w - 58, slot_h], fill=255)
        mask_draw.rectangle([slot_w - 58, 0, slot_w, 42], fill=0)
        mask = mask.filter(ImageFilter.GaussianBlur(6))
        out.paste(patch, (x1, y1), mask)
        return out.convert("RGB")


def _prepare_header_source():
    src_path = _ensure_header_source()
    if not src_path:
        return None
    from PIL import Image

    img = Image.open(src_path)
    cleaned = _clean_header_logo_slot(img)
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    cleaned.save(HEADER_BAND_IMG, quality=95)
    return cleaned


def _build_logo_for_slot() -> None:
    from PIL import Image

    if not LOGO_3D_IMG.exists():
        return

    logo = _remove_dark_background(Image.open(LOGO_3D_IMG))
    logo = _trim_transparent(logo)
    logo.thumbnail((300, 155), Image.LANCZOS)
    logo = _add_logo_shadow(logo, offset=(6, 8), blur_radius=10, opacity=55)
    logo.save(LOGO_BADGE_IMG)


def _build_full_header_image() -> None:
    from PIL import Image

    band = _prepare_header_source()
    if band is None:
        return

    _build_logo_for_slot()
    src_w, src_h = band.size
    render_h = int(src_h * HEADER_RENDER_WIDTH / src_w)
    canvas = band.resize((HEADER_RENDER_WIDTH, render_h), Image.LANCZOS).convert("RGBA")

    if LOGO_BADGE_IMG.exists():
        logo = Image.open(LOGO_BADGE_IMG).convert("RGBA")
        x1, y1, x2, y2 = LOGO_SLOT
        scale_x = HEADER_RENDER_WIDTH / src_w
        scale_y = render_h / src_h
        slot_w = int((x2 - x1) * scale_x)
        slot_h = int((y2 - y1) * scale_y)
        slot_x = int(x1 * scale_x)
        slot_y = int(y1 * scale_y)
        logo.thumbnail((slot_w - 24, slot_h - 20), Image.LANCZOS)
        paste_x = slot_x + (slot_w - logo.width) // 2
        paste_y = slot_y + (slot_h - logo.height) // 2 - 4
        canvas.paste(logo, (paste_x, paste_y), logo)

    canvas.convert("RGB").save(GOLD_HEADER_IMG, quality=95)


def _build_gold_header_image() -> None:
    if _ensure_header_source():
        _build_full_header_image()
        return

    from PIL import Image, ImageDraw

    width, height = 1654, 460
    canvas = Image.new("RGB", (width, height), f"#{WHITE}")
    draw = ImageDraw.Draw(canvas)

    margin_right = 28
    margin_top = 14
    text_left = 90

    if SOURCE_IMG.exists():
        src = Image.open(SOURCE_IMG).convert("RGB")
        sw, sh = src.size
        corner_size = 72
        for box, dest in (
            ((0, 0, corner_size, corner_size), (0, 0)),
            ((sw - corner_size, 0, sw, corner_size), (width - corner_size, 0)),
        ):
            corner = src.crop(box).resize((corner_size, corner_size), Image.LANCZOS)
            canvas.paste(corner, dest)

    _build_logo_badge()

    badge_w = 0
    if LOGO_BADGE_IMG.exists():
        badge = Image.open(LOGO_BADGE_IMG).convert("RGBA")
        badge_w = badge.width
        badge_x = width - badge.width - margin_right
        badge_y = margin_top
        canvas.paste(badge, (badge_x, badge_y), badge)

    text_box = width - badge_w - margin_right - text_left - 40
    text_bottom = margin_top + 20

    gold = f"#{GOLD_DARK}"
    gold_mid = f"#{GOLD_HEX}"
    gold_light = f"#{GOLD_LIGHT}"

    if HEADER_BAND_IMG.exists():
        band = Image.open(HEADER_BAND_IMG).convert("RGBA")
        max_band_w = text_box + 40
        max_band_h = 300
        band.thumbnail((max_band_w, max_band_h), Image.LANCZOS)
        canvas.paste(band, (text_left - 10, margin_top + 8), band)
        text_bottom = margin_top + 8 + band.height + 16
    else:
        title_ar = _load_font(44, bold=True, arabic=True)
        legal_ar = _load_font(21, arabic=True)
        tagline_ar = _load_font(19, arabic=True)
        title_en = _load_font(14, bold=True)
        tagline_en = _load_font(12)

        y = margin_top + 24
        _draw_right_in_box(draw, _shape_arabic(ORG["name_ar"]), y, text_left, text_box, title_ar, gold)
        y += 54
        draw.line([(text_left, y), (text_left + text_box - 20, y)], fill=gold_mid, width=2)
        draw.polygon(
            [(text_left + text_box - 20, y - 5), (text_left + text_box - 14, y + 2), (text_left + text_box - 20, y + 9), (text_left + text_box - 26, y + 2)],
            fill=gold_mid,
        )
        y += 14
        _draw_right_in_box(draw, _shape_arabic(ORG["name_ar_legal"]), y, text_left, text_box, legal_ar, gold_mid)
        y += 28
        y = _draw_wrapped_en_left(
            draw,
            ORG["name_en"].upper(),
            y,
            text_left,
            text_box,
            title_en,
            gold,
            line_height=19,
        )
        y += 4
        _draw_left_in_box(draw, ORG["tagline_en"].upper(), y, text_left, text_box, tagline_en, gold_mid)
        y += 20
        _draw_right_in_box(draw, _shape_arabic(ORG["tagline_ar"]), y, text_left, text_box, tagline_ar, gold_mid)
        text_bottom = y + 36

    if badge_w:
        divider_x = width - badge_w - margin_right - 18
        draw.line([(divider_x, margin_top + 10), (divider_x, text_bottom)], fill=gold_light, width=1)
        draw.line([(divider_x + 2, margin_top + 10), (divider_x + 2, text_bottom)], fill=gold_mid, width=1)

    sep_y = max(text_bottom, margin_top + 190)
    draw.line([(text_left, sep_y), (width - badge_w - margin_right - 30, sep_y)], fill=gold_light, width=1)
    draw.line([(text_left, sep_y + 2), (width - badge_w - margin_right - 30, sep_y + 2)], fill=gold_mid, width=2)

    small = _load_font(14)
    bar_y = height - 88
    draw.line([(90, bar_y), (width - 90, bar_y)], fill=gold_mid, width=2)
    draw.line([(90, bar_y + 3), (width - 90, bar_y + 3)], fill=gold_light, width=1)

    loc = _shape_arabic(f"الموقع: {ORG['location_ar']}")
    phones = f"{ORG['phones'][0]}  ·  {ORG['phones'][1]}"
    cr_line = f"C.R. {ORG['cr']}  ·  P.O. Box {ORG['postal']}  ·  {ORG['email']}"

    _draw_centered_in_box(draw, loc, bar_y + 14, 0, width, small, gold)
    _draw_centered_in_box(draw, phones, bar_y + 36, 0, width, small, gold)
    _draw_centered_in_box(draw, cr_line, bar_y + 58, 0, width, small, gold_mid)

    canvas.save(GOLD_HEADER_IMG)


def create_gold_letterhead_assets() -> None:
    if not SOURCE_IMG.exists():
        print(f"Error: letterhead source image not found at {SOURCE_IMG}", file=sys.stderr)
        sys.exit(1)

    try:
        from PIL import Image
    except ImportError:
        print("Error: Pillow is required. Run: pip install pillow", file=sys.stderr)
        sys.exit(1)

    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    if uses_full_header():
        _build_full_header_image()
    else:
        _build_logo_badge()
        _build_gold_header_image()

    if uses_full_header():
        return

    footer_src = DOCS_DIR / "letterhead-gold-footer.png"
    if footer_src.exists() and footer_src.stat().st_size > 10000:
        from PIL import Image

        footer = Image.open(footer_src).convert("RGB")
        fw, fh = footer.size
        footer = footer.resize((1654, int(1654 * fh / fw)), Image.LANCZOS)
        footer.save(GOLD_FOOTER_IMG)
        return

    src = Image.open(SOURCE_IMG).convert("RGB")
    w, h = src.size
    footer_h = int(h * 0.075)
    footer = src.crop((0, h - footer_h, w, h))
    footer = footer.resize((1654, int(1654 * footer_h / w)), Image.LANCZOS)
    footer.save(GOLD_FOOTER_IMG)


def add_full_width_image(doc, image_path: Path, width_cm: float) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))


def create_gold_base_document() -> Document:
    create_gold_letterhead_assets()

    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(0.5) if uses_full_header() else Cm(0.9)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.0) if uses_full_header() else Cm(1.4)
    section.right_margin = Cm(1.0) if uses_full_header() else Cm(1.4)
    if not uses_full_header():
        set_page_borders(section)

    if GOLD_HEADER_IMG.exists():
        add_full_width_image(doc, GOLD_HEADER_IMG, 18.2 if uses_full_header() else 17.8)

    if not uses_full_header():
        add_gold_contact_line(doc)

    return doc


def add_gold_contact_line(doc) -> None:
    for line, rtl in ((CONTACT_LINE_EN, False), (CONTACT_LINE_AR, True)):
        p = doc.add_paragraph()
        if rtl:
            set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        set_run_font(run, size=8, color=MUTED)


def add_gold_divider(doc) -> None:
    table = doc.add_table(rows=1, cols=1)
    remove_table_borders(table)
    cell = table.rows[0].cells[0]
    set_cell_margins(cell, top=40, start=0, bottom=40, end=0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("◆")
    set_run_font(run, size=10, color=GOLD_RGB)


def add_gold_content_area(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)
    cell = table.rows[0].cells[0]
    set_cell_margins(cell, top=60, start=120, bottom=60, end=120)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    return cell


def add_gold_signature_block(doc) -> None:
    table = doc.add_table(rows=1, cols=2)
    remove_table_borders(table)
    table.columns[0].width = Cm(8)
    table.columns[1].width = Cm(8)

    ar_cell = table.rows[0].cells[0]
    en_cell = table.rows[0].cells[1]
    set_cell_margins(ar_cell, top=160, start=80, bottom=40, end=80)
    set_cell_margins(en_cell, top=160, start=80, bottom=40, end=80)

    ar_p = ar_cell.paragraphs[0]
    set_rtl(ar_p)
    ar_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for text, bold in [
        ("مع التقدير،", False),
        ("_________________________", False),
        ("الاسم والتوقيع", True),
        (ORG["name_ar_legal"], False),
    ]:
        run = ar_p.add_run(text + "\n")
        set_run_font(run, size=10, bold=bold, color=GOLD_TEXT)

    en_p = en_cell.paragraphs[0]
    en_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for text, bold in [
        ("Regards,", False),
        ("_________________________", False),
        ("Name & Signature", True),
        (ORG["name_en"], False),
    ]:
        run = en_p.add_run(text + "\n")
        set_run_font(run, size=10, bold=bold, color=GOLD_TEXT)


def add_gold_footer(doc) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(FOOTER_LINE)
    set_run_font(run, size=8, color=MUTED)

    if GOLD_FOOTER_IMG.exists():
        add_full_width_image(doc, GOLD_FOOTER_IMG, 17.8)


def finalize_gold_document(doc: Document) -> Document:
    if not uses_full_header():
        add_gold_footer(doc)
    return doc
