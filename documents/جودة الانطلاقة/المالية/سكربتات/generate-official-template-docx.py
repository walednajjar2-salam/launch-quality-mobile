#!/usr/bin/env python3
"""Generate gold-bordered official Word letterhead template (.docx)."""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx.enum.text import WD_ALIGN_PARAGRAPH

from letterhead_gold_docx import (
    DOCS_DIR,
    DARK,
    GOLD_RGB,
    GOLD_TEXT,
    MUTED,
    add_gold_content_area,
    add_gold_divider,
    add_gold_signature_block,
    create_gold_base_document,
    finalize_gold_document,
    set_rtl,
    set_run_font,
)
from org_info import ORG

OUTPUT = DOCS_DIR / "قالب-رسمي.docx"

TEMPLATE_LINES = [
    (f"الرقم: {ORG['ref_prefix']} / ____ / ____", 10, False, MUTED, 8),
    ("التاريخ: ____ / ____ / ______ هـ  ·  ____ / ____ / ______ م", 10, False, MUTED, 10),
    ("", 11, False, MUTED, 4),
    ("إلى: _________________________________________________", 11, False, GOLD_TEXT, 8),
    ("الموضوع: _____________________________________________", 11, True, GOLD_RGB, 14),
    ("", 12, False, MUTED, 6),
    ("السلام عليكم ورحمة الله وبركاته،", 13, False, DARK, 12),
    ("", 12, False, MUTED, 6),
    ("[اكتب نص الخطاب هنا — يمكنك حذف هذا السطر والكتابة مباشرة]", 12, False, MUTED, 10),
    ("", 12, False, MUTED, 8),
    ("", 12, False, MUTED, 8),
    ("", 12, False, MUTED, 8),
    ("", 12, False, MUTED, 8),
    ("وتفضلوا بقبول فائق الاحترام والتقدير،", 12, False, DARK, 6),
]


def build_template():
    doc = create_gold_base_document()
    add_gold_divider(doc)
    content_cell = add_gold_content_area(doc)

    first = True
    for text, size, bold, color, spacing_after in TEMPLATE_LINES:
        if first:
            p = content_cell.paragraphs[0]
            first = False
        else:
            p = content_cell.add_paragraph()
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.line_spacing = 1.85
        p.paragraph_format.space_after = spacing_after
        if text:
            run = p.add_run(text)
            set_run_font(run, size=size, bold=bold, color=color)

    add_gold_signature_block(doc)
    return finalize_gold_document(doc)


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_template()
    try:
        doc.save(OUTPUT)
        print(f"Created: {OUTPUT}")
    except PermissionError:
        fallback = DOCS_DIR / "official-template-latest.docx"
        doc.save(fallback)
        print(f"Warning: close {OUTPUT.name} in Word, then re-run.", file=sys.stderr)
        print(f"Created: {fallback}")


if __name__ == "__main__":
    main()
