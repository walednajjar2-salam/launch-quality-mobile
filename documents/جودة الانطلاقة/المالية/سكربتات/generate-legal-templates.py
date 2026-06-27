#!/usr/bin/env python3
"""Generate gold-letterhead contract, invoice, and official letter templates."""

from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx.enum.text import WD_ALIGN_PARAGRAPH

from letterhead_gold_docx import (
    DARK,
    GOLD_RGB,
    GOLD_TEXT,
    MUTED,
    add_gold_content_area,
    add_gold_divider,
    add_gold_signature_block,
    create_gold_base_document,
    finalize_gold_document,
    set_rtl,
    set_run_font,
)
from org_info import ORG

RAILWAY_DOCS = Path(r"C:\Users\walee\Projects\Launch-Quality-Railway-Deploy\public\documents")
LOCAL_DOCS = Path(__file__).resolve().parent.parent / "documents"
EMPLOY_DOCS = Path(r"d:\qulaty of lunch\employ")


def add_block(cell, text, *, size=11, bold=False, color=DARK, space_after=8):
    p = cell.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.75
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def build_contract_template():
    doc = create_gold_base_document()
    add_gold_divider(doc)
    cell = add_gold_content_area(doc)
    p = cell.paragraphs[0]
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("عقد إيجار محمي · Protected Lease Agreement")
    set_run_font(run, size=18, bold=True, color=GOLD_RGB)
    add_block(cell, f"الرقم: {ORG['ref_prefix']}-LEASE-____ / ____", size=10, color=MUTED)
    add_block(cell, "التاريخ: ____ / ____ / ______", size=10, color=MUTED, space_after=14)
    add_block(cell, "الطرف الأول (المؤجر): " + ORG["name_ar_legal"], bold=True, color=GOLD_TEXT)
    add_block(cell, "السجل التجاري: " + ORG["cr"] + " · " + ORG["location_ar"])
    add_block(cell, "الطرف الثاني (المستأجر): _________________________________")
    add_block(cell, "الهوية / السجل: __________________ · الجنسية: ________________")
    add_block(cell, "1) العين المؤجرة", bold=True, color=GOLD_RGB, space_after=4)
    add_block(cell, "المبنى: ____ · الشقة/الوحدة: ____ · الغرفة: ____ · الموقع: ________________")
    add_block(cell, "2) مدة العقد", bold=True, color=GOLD_RGB, space_after=4)
    add_block(cell, "من ____ / ____ / ______ إلى ____ / ____ / ______ · دورة الدفع: شهرية")
    add_block(cell, "3) القيمة المالية", bold=True, color=GOLD_RGB, space_after=4)
    add_block(cell, "الإيجار الشهري: ______ OMR · التأمين: ______ OMR · غرامة التأخير: ______ OMR")
    add_block(cell, "4) البنود القانونية", bold=True, color=GOLD_RGB, space_after=4)
    add_block(cell, "[أدخل البنود المعتمدة: التجديد، الإخلاء، الصيانة، الغرامات، التوقيع والختم]")
    add_block(cell, "5) التوقيعات", bold=True, color=GOLD_RGB, space_after=6)
    add_block(cell, "المؤجر: _________________________    المستأجر: _________________________")
    add_block(cell, "الختم: __________________________    التاريخ: ____ / ____ / ______")
    add_gold_signature_block(doc)
    return finalize_gold_document(doc)


def build_invoice_template():
    doc = create_gold_base_document()
    add_gold_divider(doc)
    cell = add_gold_content_area(doc)
    p = cell.paragraphs[0]
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("فاتورة · Tax Invoice")
    set_run_font(run, size=18, bold=True, color=GOLD_RGB)
    add_block(cell, "رقم الفاتورة: INV-____-________", size=10, color=MUTED)
    add_block(cell, "تاريخ الإصدار: ____ / ____ / ______ · الاستحقاق: ____ / ____ / ______", size=10, color=MUTED, space_after=14)
    add_block(cell, ORG["name_ar_legal"], bold=True, color=GOLD_TEXT)
    add_block(cell, ORG["email"] + " · " + ORG["phones"][0])
    add_block(cell, "العميل: _________________________________")
    add_block(cell, "العقار / العقد: _________________________ · REF: _______________", space_after=12)
    add_block(cell, "البيان                    الكمية        السعر        الإجمالي", bold=True)
    add_block(cell, "______________________________________________________________")
    add_block(cell, "______________________________________________________________")
    add_block(cell, "المجموع: __________ OMR · المدفوع: __________ · المتبقي: __________", bold=True, color=GOLD_TEXT)
    add_block(cell, "طريقة الدفع: □ نقدي  □ تحويل بنكي  □ شيك  · المرجع: ______________")
    add_gold_signature_block(doc)
    return finalize_gold_document(doc)


def build_official_letter():
    doc = create_gold_base_document()
    add_gold_divider(doc)
    cell = add_gold_content_area(doc)
    add_block(cell, f"الرقم: {ORG['ref_prefix']}-LTR-____ / ____", size=10, color=MUTED)
    add_block(cell, "التاريخ: ____ / ____ / ______ هـ", size=10, color=MUTED, space_after=10)
    add_block(cell, "إلى: _________________________________________________")
    add_block(cell, "الموضوع: _____________________________________________", bold=True, color=GOLD_RGB, space_after=12)
    add_block(cell, "السلام عليكم ورحمة الله وبركاته،")
    add_block(cell, "[اكتب نص الخطاب الرسمي هنا]", color=MUTED)
    add_block(cell, "", space_after=20)
    add_block(cell, "وتفضلوا بقبول فائق الاحترام والتقدير،")
    add_gold_signature_block(doc)
    return finalize_gold_document(doc)


TEMPLATES = [
    ("lq-contract-template.docx", build_contract_template),
    ("lq-invoice-template.docx", build_invoice_template),
    ("lq-official-letter.docx", build_official_letter),
]


def save_all():
    for folder in (RAILWAY_DOCS, LOCAL_DOCS, EMPLOY_DOCS):
        folder.mkdir(parents=True, exist_ok=True)
    for name, builder in TEMPLATES:
        doc = builder()
        for folder in (RAILWAY_DOCS, LOCAL_DOCS, EMPLOY_DOCS):
            out = folder / name
            doc.save(out)
            print(f"Created: {out}")


if __name__ == "__main__":
    save_all()
