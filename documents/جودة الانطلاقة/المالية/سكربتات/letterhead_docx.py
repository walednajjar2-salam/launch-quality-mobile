"""Shared letterhead helpers for QUALITY OF LAUNCH PROJECTS LLC Word documents."""

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from org_info import ORG

REPO_ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = REPO_ROOT / "documents"
HEADER_IMG = DOCS_DIR / "letterhead-header.png"
FOOTER_IMG = DOCS_DIR / "letterhead-footer.png"
LOGO_SRC = REPO_ROOT / "assets" / "logo.png"
LOGO_IMG = DOCS_DIR / "logo.png"

NAVY = "0D1728"
TEAL = "05AAA1"
GOLD = "C9A227"
CREAM = "F7EEDC"
DARK = RGBColor(0x11, 0x18, 0x27)
MUTED = RGBColor(0x6B, 0x72, 0x80)
NAVY_RGB = RGBColor(0x0D, 0x17, 0x28)


def set_rtl(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    p_pr.append(bidi)


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tc_pr.append(margins)


def remove_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    tbl_pr.append(borders)


def set_run_font(run, size=11, bold=False, color=DARK, name="Arial") -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:cs"), name)


def set_page_background(doc, color_hex: str) -> None:
    background = OxmlElement("w:background")
    background.set(qn("w:color"), color_hex)
    doc.element.insert(0, background)


def load_font(size: int, bold: bool = False):
    from PIL import ImageFont

    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf") if bold else Path("C:/Windows/Fonts/arial.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")
        if bold
        else Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def sync_logo_copy() -> None:
    if not LOGO_SRC.exists():
        print(f"Error: official logo not found at {LOGO_SRC}", file=sys.stderr)
        print("Place the official logo at assets/logo.png and re-run.", file=sys.stderr)
        sys.exit(1)

    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(LOGO_SRC, LOGO_IMG)


def paste_logo_on_header(header) -> None:
    from PIL import Image

    logo = Image.open(LOGO_SRC).convert("RGBA")
    w, h = header.size

    logo_area_left = int(w * 0.52)
    logo_area_right = w - 20
    logo_area_top = 10
    logo_area_bottom = h - 10
    max_w = logo_area_right - logo_area_left
    max_h = logo_area_bottom - logo_area_top

    logo.thumbnail((max_w, max_h), Image.LANCZOS)
    x = logo_area_left + (max_w - logo.width) // 2
    y = logo_area_top + (max_h - logo.height) // 2
    header.paste(logo, (x, y), logo)


def create_letterhead_images() -> None:
    try:
        from PIL import Image, ImageDraw
    except ImportError:
        return

    sync_logo_copy()

    w, h = 1240, 220
    header = Image.new("RGB", (w, h), f"#{CREAM}")
    draw = ImageDraw.Draw(header)

    navy = [(int(w * 0.48), 0), (w, 0), (w, int(h * 0.92)), (0, h)]
    draw.polygon(navy, fill=f"#{NAVY}")
    draw.polygon([(0, int(h * 0.64)), (int(w * 0.19), int(h * 0.64)), (0, h)], fill=f"#{TEAL}")
    draw.polygon(
        [(int(w * 0.73), 0), (int(w * 0.84), 0), (int(w * 0.79), int(h * 0.34)), (int(w * 0.68), int(h * 0.34))],
        fill=f"#{GOLD}",
    )
    draw.polygon(
        [(int(w * 0.66), 0), (int(w * 0.81), 0), (int(w * 0.76), int(h * 0.47)), (int(w * 0.61), int(h * 0.47))],
        fill=f"#{TEAL}",
    )

    title_font = load_font(34, bold=True)
    sub_font = load_font(14, bold=True)

    en_lines = ORG["name_en"].upper().replace(" & ", " &\n", 1).replace(" LLC", "\nLLC")
    draw.text((48, 36), en_lines, fill="#111827", font=title_font, spacing=4)
    draw.text((48, 138), ORG["tagline_en"].upper(), fill=f"#{GOLD}", font=sub_font)

    paste_logo_on_header(header)
    header.save(HEADER_IMG)

    fw, fh = 1240, 280
    footer = Image.new("RGB", (fw, fh), f"#{CREAM}")
    draw = ImageDraw.Draw(footer)
    draw.rectangle([(0, fh - 56), (fw, fh)], fill=f"#{NAVY}")
    draw.polygon([(int(fw * 0.28), 0), (fw, 0), (fw, fh - 56), (0, fh - 56)], fill=f"#{TEAL}")
    for x_offset, color, width in ((455, GOLD, 35), (420, TEAL, 20), (395, CREAM, 8)):
        draw.polygon(
            [
                (fw - x_offset, 0),
                (fw - x_offset + width, 0),
                (fw - x_offset + width - 18, fh - 56),
                (fw - x_offset - 18, fh - 56),
            ],
            fill=f"#{color}",
        )

    contact_font = sub_font
    contacts = [
        (ORG["email"], 58),
        (ORG["phones"][1], 98),
        (ORG["website"], 138),
        (ORG["location_en"], 178),
    ]
    for text, y in contacts:
        draw.text((fw - 420, y), text, fill="white", font=contact_font)

    footer.save(FOOTER_IMG)


def add_full_width_image(doc, image_path: Path, width_cm: float) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))


def create_base_document() -> Document:
    create_letterhead_images()

    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(0)
    section.bottom_margin = Cm(0)
    section.left_margin = Cm(0)
    section.right_margin = Cm(0)

    set_page_background(doc, CREAM)

    if HEADER_IMG.exists():
        add_full_width_image(doc, HEADER_IMG, 21.0)

    return doc


def add_content_box(doc: Document):
    content_table = doc.add_table(rows=1, cols=1)
    content_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(content_table)
    content_cell = content_table.rows[0].cells[0]
    set_cell_shading(content_cell, "FFFFFF")
    set_cell_margins(content_cell, top=180, start=360, bottom=180, end=360)

    tc_pr = content_cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "E8DFC8")
        borders.append(element)
    tc_pr.append(borders)

    return content_cell


def add_rtl_paragraph(cell, text: str, *, size=12, bold=False, color=DARK, align=WD_ALIGN_PARAGRAPH.RIGHT, spacing_after=6):
    p = cell.paragraphs[0] if not cell.paragraphs[0].text and len(cell.paragraphs) == 1 else cell.add_paragraph()
    if cell.paragraphs[0] is p and p.text:
        p = cell.add_paragraph()
    set_rtl(p)
    p.alignment = align
    p.paragraph_format.line_spacing = 1.9
    p.paragraph_format.space_after = Pt(spacing_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_signature_block(doc) -> None:
    sig_table = doc.add_table(rows=1, cols=1)
    remove_table_borders(sig_table)
    sig_cell = sig_table.rows[0].cells[0]
    set_cell_margins(sig_cell, top=120, start=360, bottom=60, end=360)

    sig_p = sig_cell.paragraphs[0]
    sig_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for i, (text, bold) in enumerate(
        [
            ("Regards,", False),
            ("Signature", True),
            (ORG["name_en"], False),
        ]
    ):
        run = sig_p.add_run(text + ("\n" if i < 2 else ""))
        set_run_font(run, size=10, bold=bold, color=DARK)


def finalize_document(doc: Document) -> Document:
    if FOOTER_IMG.exists():
        add_full_width_image(doc, FOOTER_IMG, 21.0)
    return doc
