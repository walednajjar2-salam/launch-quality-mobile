#!/usr/bin/env python3
"""Generate premium gold-letterhead employee responsibilities document (.docx)."""

from __future__ import annotations

import re
import sys
from dataclasses import dataclass, field
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from letterhead_gold_docx import (
    DARK,
    GOLD_HEX,
    GOLD_RGB,
    GOLD_TEXT,
    MUTED,
    add_gold_content_area,
    add_gold_divider,
    create_gold_base_document,
    finalize_gold_document,
    remove_table_borders,
    set_cell_margins,
    set_cell_shading,
    set_rtl,
    set_run_font,
)
from org_info import ORG as ORG_INFO

SOURCE_TXT = Path(r"d:\qulaty of lunch\employ\employ_extract.txt")
FALLBACK_TXT = Path(__file__).resolve().parent.parent / "employ_extract.txt"
OUTPUT = Path(r"d:\qulaty of lunch\employ\Emplooy - Premium.docx")
OUTPUT_FALLBACK = Path(__file__).resolve().parent.parent / "documents" / "Emplooy-Premium.docx"

SECTION_MARK = re.compile(r"^(أولاً|ثانياً|ثالثاً|رابعاً|خامساً|سادساً|سابعاً|ثامناً|تاسعاً|عاشراً)[：:]?\s*(.+)$")
META_MARK = re.compile(r"^(الاسم|الجنسية|الرقم المدني|المسمى الوظيفي|الاسم|التوقيع)[：:]\s*(.*)$")


@dataclass
class RoleBlock:
    ordinal: str
    title: str
    name: str = ""
    nationality: str = ""
    civil_id: str = ""
    duties: list[str] = field(default_factory=list)


@dataclass
class Signatory:
    title: str
    name: str


def load_lines() -> list[str]:
    path = SOURCE_TXT if SOURCE_TXT.exists() else FALLBACK_TXT
    return [ln.strip() for ln in path.read_text(encoding="utf-8").splitlines() if ln.strip()]


def parse_document(lines: list[str]) -> tuple[str, list[RoleBlock], list[str], list[Signatory], list[str]]:
    title = lines[0] if lines else "الصلاحيات والمسؤوليات الوظيفية"
    roles: list[RoleBlock] = []
    org_lines: list[str] = []
    signatories: list[Signatory] = []
    closing: list[str] = []

    current: RoleBlock | None = None
    mode = "roles"

    for line in lines[1:]:
        if line == "التسلسل الإداري":
            if current:
                roles.append(current)
                current = None
            mode = "org"
            continue
        if line == "إقرار واعتماد المسؤوليات الوظيفية":
            mode = "ack"
            continue
        if line == "اعتماد اللائحة":
            mode = "approval"
            continue

        if mode == "roles":
            m = SECTION_MARK.match(line)
            if m:
                if current:
                    roles.append(current)
                current = RoleBlock(ordinal=m.group(1), title=m.group(2).strip())
                continue
            meta = META_MARK.match(line)
            if meta and current:
                key, val = meta.group(1), meta.group(2).strip()
                if key == "الاسم":
                    current.name = val
                elif key == "الجنسية":
                    current.nationality = val
                elif key == "الرقم المدني":
                    current.civil_id = val
                continue
            if line == "الصلاحيات والمسؤوليات:":
                continue
            if current and line:
                current.duties.append(line)
            continue

        if mode == "org":
            if line.startswith("↓"):
                org_lines.append("↓")
            else:
                org_lines.append(line)
            continue

        if mode == "ack":
            if line in {"المسمى الوظيفي", "الاسم", "التوقيع"}:
                continue
            if line.startswith("_"):
                if signatories:
                    signatories[-1].name = signatories[-1].name or "__________________"
                continue
            if signatories and not signatories[-1].name and not line.startswith("_"):
                if any(ch.isalpha() for ch in line):
                    signatories[-1].name = line
                    continue
            if line in {
                "المدير العام",
                "منسق المدير العام والشؤون الإدارية للموظفين",
                "موظفة الاستقبال",
                "مسؤولة قسم التسويق والتصميم",
                "المحاسب",
                "مدير إدارة الضيافة وتنظيم الحفلات",
                "مدير إدارة العلاقات والعقارات",
            }:
                signatories.append(Signatory(title=line, name=""))
                continue
            if not signatories and "يقر" in line:
                closing.append(line)
            elif not signatories:
                closing.append(line)
            continue

        if mode == "approval":
            closing.append(line)

    if current:
        roles.append(current)
    return title, roles, org_lines, signatories, closing


def add_paragraph(cell, text, *, size=11, bold=False, color=DARK, align=WD_ALIGN_PARAGRAPH.RIGHT, rtl=True, space_after=6, line_spacing=1.75):
    p = cell.add_paragraph()
    if rtl:
        set_rtl(p)
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_gold_heading_bar(cell, text: str) -> None:
    table = cell.add_table(rows=1, cols=1)
    remove_table_borders(table)
    c = table.rows[0].cells[0]
    set_cell_shading(c, GOLD_HEX)
    set_cell_margins(c, top=80, start=120, bottom=80, end=120)
    p = c.paragraphs[0]
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))


def add_role_card(cell, role: RoleBlock, index: int) -> None:
    card = cell.add_table(rows=1, cols=1)
    remove_table_borders(card)
    box = card.rows[0].cells[0]
    set_cell_margins(box, top=100, start=140, bottom=100, end=140)
    set_cell_shading(box, "FBF7EE")

    tc_pr = box._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "12")
        b.set(qn("w:color"), GOLD_HEX)
        borders.append(b)
    tc_pr.append(borders)

    header = box.paragraphs[0]
    set_rtl(header)
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(4)
    r1 = header.add_run(f"{role.ordinal} — {role.title}")
    set_run_font(r1, size=13, bold=True, color=GOLD_RGB)

    meta_bits = []
    if role.name and not role.name.startswith("_"):
        meta_bits.append(f"الاسم: {role.name}")
    if role.nationality:
        meta_bits.append(f"الجنسية: {role.nationality}")
    if role.civil_id and not role.civil_id.startswith("_"):
        meta_bits.append(f"الرقم المدني: {role.civil_id}")
    if meta_bits:
        mp = box.add_paragraph()
        set_rtl(mp)
        mp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        mp.paragraph_format.space_after = Pt(8)
        run = mp.add_run("   ·   ".join(meta_bits))
        set_run_font(run, size=10, bold=False, color=GOLD_TEXT)

    sub = box.add_paragraph()
    set_rtl(sub)
    sub.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sub.paragraph_format.space_after = Pt(6)
    run = sub.add_run("الصلاحيات والمسؤوليات")
    set_run_font(run, size=11, bold=True, color=DARK)

    for i, duty in enumerate(role.duties, 1):
        p = box.add_paragraph()
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.left_indent = Cm(0.15)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.65
        marker = p.add_run(f"{i}. ")
        set_run_font(marker, size=10, bold=True, color=GOLD_RGB)
        body = p.add_run(duty)
        set_run_font(body, size=10.5, color=DARK)

    spacer = cell.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)


def add_org_chart(cell, org_lines: list[str]) -> None:
    add_gold_heading_bar(cell, "التسلسل الإداري")
    add_paragraph(cell, "", space_after=8)

    for line in org_lines:
        if line == "↓":
            add_paragraph(cell, "▼", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, color=GOLD_RGB, space_after=2)
            continue
        table = cell.add_table(rows=1, cols=1)
        remove_table_borders(table)
        c = table.rows[0].cells[0]
        set_cell_shading(c, "FFFDF8")
        set_cell_margins(c, top=60, start=200, bottom=60, end=200)
        tc_pr = c._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{edge}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "8")
            b.set(qn("w:color"), "D4AF37")
            borders.append(b)
        tc_pr.append(borders)
        p = c.paragraphs[0]
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        set_run_font(run, size=11, bold=True, color=GOLD_TEXT)
        add_paragraph(cell, "", space_after=4)


def add_signatures_table(cell, signatories: list[Signatory]) -> None:
    add_gold_heading_bar(cell, "إقرار واعتماد المسؤوليات الوظيفية")
    add_paragraph(
        cell,
        "يقر كل موظف بالاطلاع على مهامه وصلاحياته الواردة في هذه اللائحة، "
        "والالتزام بتنفيذها وفق الأنظمة والسياسات المعتمدة بالشركة.",
        size=11,
        color=DARK,
        space_after=12,
    )

    table = cell.add_table(rows=1, cols=3)
    remove_table_borders(table)
    headers = ["المسمى الوظيفي", "الاسم", "التوقيع"]
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        set_cell_shading(c, GOLD_HEX)
        set_cell_margins(c, top=60, start=60, bottom=60, end=60)
        p = c.paragraphs[0]
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    for sig in signatories:
        row = table.add_row().cells
        values = [sig.title, sig.name or "__________________", "__________________"]
        for i, val in enumerate(values):
            set_cell_margins(row[i], top=80, start=60, bottom=80, end=60)
            if i == 0:
                set_cell_shading(row[i], "FBF7EE")
            p = row[i].paragraphs[0]
            set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i else WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(val)
            set_run_font(run, size=10, color=DARK if i else GOLD_TEXT, bold=(i == 0))

    add_paragraph(cell, "", space_after=10)


def add_approval_block(cell, closing: list[str]) -> None:
    add_gold_heading_bar(cell, "اعتماد اللائحة")
    for line in closing:
        if line.startswith("اسم الشركة") or line.startswith("التاريخ") or line.startswith("ختم") or line.startswith("توقيع"):
            add_paragraph(cell, line, size=11, bold=True, color=GOLD_TEXT, space_after=8)
        elif line.startswith("الاسم:"):
            add_paragraph(cell, line, size=11, color=DARK, space_after=14)
        else:
            add_paragraph(cell, line, size=11, color=DARK, space_after=8)


def build_document() -> "Document":
    lines = load_lines()
    title, roles, org_lines, signatories, closing = parse_document(lines)

    doc = create_gold_base_document()
    add_gold_divider(doc)
    content = add_gold_content_area(doc)

    # Title block
    p = content.paragraphs[0]
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    set_run_font(run, size=20, bold=True, color=GOLD_RGB)

    sub = content.add_paragraph()
    set_rtl(sub)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(6)
    run = sub.add_run(ORG_INFO["name_ar_legal"])
    set_run_font(run, size=11, color=MUTED)

    meta = content.add_paragraph()
    set_rtl(meta)
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(14)
    run = meta.add_run(f"الإصدار: 2026   ·   المرجع: {ORG_INFO['ref_prefix']}-HR-001")
    set_run_font(run, size=9, color=MUTED)

    add_gold_divider(doc)

    for role in roles:
        add_role_card(content, role, 0)

    if org_lines:
        add_org_chart(content, org_lines)
    if signatories:
        add_signatures_table(content, signatories)
    if closing:
        add_approval_block(content, closing)

    return finalize_gold_document(doc)


def main() -> None:
    doc = build_document()
    targets = [OUTPUT, OUTPUT_FALLBACK]
    saved = None
    for target in targets:
        try:
            target.parent.mkdir(parents=True, exist_ok=True)
            doc.save(target)
            saved = target
            print(f"Created: {target}")
        except PermissionError:
            continue
    if not saved:
        alt = OUTPUT.with_name("Emplooy-Premium-latest.docx")
        doc.save(alt)
        print(f"Created: {alt}")
        print("Tip: close the original file in Word if it is open.", file=sys.stderr)


if __name__ == "__main__":
    main()
