#!/usr/bin/env python3
"""Generate QUALITY OF LAUNCH PROJECTS LLC hospitality manager letter (.docx)."""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx.enum.text import WD_ALIGN_PARAGRAPH

from letterhead_docx import (
    DARK,
    DOCS_DIR,
    NAVY_RGB,
    add_content_box,
    add_signature_block,
    create_base_document,
    finalize_document,
    set_rtl,
    set_run_font,
)

OUTPUT = DOCS_DIR / "مدير-إدارة-الضيافة.docx"

RESPONSIBILITIES = [
    "إدارة الحجوزات والفعاليات",
    "متابعة تجهيز الحفلات والمناسبات",
    "الإشراف على فرق التشغيل والعملاء",
    "إدارة المعدات والعروض الخاصة بالحفلات والفعاليات",
    "متابعة جودة الخدمة المقدمة بالفعاليات",
    "استلام وتسليم معدات وتجهيزات الحفلات",
    "متابعة كميات وأصناف المعدات والتجهيزات",
    "التأكد من سلامة وصلاحية المعدات",
    "مطابقة الموجودات والتالف ورفع طلبات الشراء",
    "المحافظة على جاهزية تنظيم الفعالية مع سجلات المخزون",
    "متابعة نقل المعدات لأي فعالية",
    "التنسيق مع قسم النقليات لتوفير وسائل النقل اللازمة للفعاليات والفرق التشغيلية",
]


def build_document():
    doc = create_base_document()
    content_cell = add_content_box(doc)

    job_p = content_cell.paragraphs[0]
    set_rtl(job_p)
    job_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = job_p.add_run("مدير إدارة الضيافة وتنظيم الحفلات")
    set_run_font(run, size=16, bold=True, color=NAVY_RGB)
    job_p.paragraph_format.space_after = 14

    for item in RESPONSIBILITIES:
        p = content_cell.add_paragraph(style="List Bullet")
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.line_spacing = 1.9
        p.paragraph_format.space_after = 3
        run = p.add_run(item)
        set_run_font(run, size=12, color=DARK)

    add_signature_block(doc)
    return finalize_document(doc)


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
