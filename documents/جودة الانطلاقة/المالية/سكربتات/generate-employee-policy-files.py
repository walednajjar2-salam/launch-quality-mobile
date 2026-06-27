#!/usr/bin/env python3
"""Generate per-employee policy files from owner-approved Templet.docx."""

from __future__ import annotations

import re
import shutil
import sys
from dataclasses import dataclass, field
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from letterhead_gold_docx import set_rtl, set_run_font
from org_info import ORG

GOLD = RGBColor(0xB8, 0x86, 0x0B)
GOLD_TEXT = RGBColor(0x7A, 0x5C, 0x10)
DARK = RGBColor(0x2D, 0x22, 0x12)
MUTED = RGBColor(0x9C, 0x8A, 0x6E)

ROLE_HEADER = re.compile(
    r"^(\d+)\)\s*(.+?)\s*[—–-]\s*المهام\s*/\s*المسؤوليات\s*/\s*الواجبات\s*$"
)
SECTION_MARK = re.compile(r"^(المهام|المسؤوليات|الواجبات)$")
BULLET = re.compile(r"^[•·\-]\s*(.+)$")


@dataclass
class RolePolicy:
    number: int
    title: str
    tasks: list[str] = field(default_factory=list)
    policies: list[str] = field(default_factory=list)
    commitments: list[str] = field(default_factory=list)


from employee_roster import EMPLOYEES, Employee, employee_files_dir, employ_dir, template_path


def duties_source_path() -> Path:
    employ = template_path().parent
    preferred = employ / "المهام والمسوليات والواجبات.docx"
    if preferred.exists():
        return preferred
    return Path(__file__).resolve().parent / "dump_المهام والمسوليات والواجبات.txt"


def output_dir() -> Path:
    return employee_files_dir()


def safe_filename(name: str, title: str) -> str:
    clean_name = re.sub(r'[\\/:*?"<>|]', "-", name.strip())
    clean_title = re.sub(r'[\\/:*?"<>|]', "-", title.strip())[:40]
    return f"{clean_name} — {clean_title}.docx"


def load_policy_lines() -> list[str]:
    src = duties_source_path()
    if src.suffix == ".docx":
        doc = Document(str(src))
        return [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return [ln.strip() for ln in src.read_text(encoding="utf-8").splitlines() if ln.strip()]


SECTION_LIMITS = {"المهام": 12, "المسؤوليات": 12, "الواجبات": 8}


def parse_role_policies(lines: list[str]) -> dict[int, RolePolicy]:
    roles: dict[int, RolePolicy] = {}
    current: RolePolicy | None = None
    section: str | None = None

    for line in lines:
        if line == "---":
            continue
        header = ROLE_HEADER.match(line)
        if header:
            num = int(header.group(1))
            if num in roles:
                break
            current = RolePolicy(num, header.group(2).strip())
            roles[current.number] = current
            section = None
            continue
        if SECTION_MARK.match(line):
            section = SECTION_MARK.match(line).group(1)
            continue
        bullet = BULLET.match(line)
        if not bullet or not current or not section:
            continue
        limit = SECTION_LIMITS.get(section, 12)
        bucket = {
            "المهام": current.tasks,
            "المسؤوليات": current.policies,
            "الواجبات": current.commitments,
        }[section]
        if len(bucket) >= limit:
            continue
        bucket.append(bullet.group(1).strip())

    return roles


def add_heading(doc: Document, text: str, size: int = 16, color=GOLD, space_before=18) -> None:
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=color)


def add_meta_line(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{label}: ")
    set_run_font(run, size=11, bold=True, color=GOLD_TEXT)
    run = p.add_run(value)
    set_run_font(run, size=11, color=DARK)


def add_section(doc: Document, title: str, items: list[str]) -> None:
    add_heading(doc, title, size=13, color=GOLD_TEXT, space_before=14)
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.left_indent = Cm(0.2)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.6
        marker = p.add_run(f"{i}. ")
        set_run_font(marker, size=10.5, bold=True, color=GOLD)
        body = p.add_run(item)
        set_run_font(body, size=10.5, color=DARK)


def add_acknowledgment(doc: Document, employee: Employee) -> None:
    add_heading(doc, "إقرار الالتزام", size=13, color=GOLD_TEXT, space_before=16)
    text = (
        f"أقر أنا {employee.name} — {employee.job_title} — "
        "بأنني اطلعت على المهام والسياسات والالتزامات الواردة أعلاه، "
        " وأتعهد بالالتزام التام بتنفيذها وفق أنظمة الشركة وتعليماتها المعتمدة."
    )
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = 1.75
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    set_run_font(run, size=11, color=DARK)

    for label in ("اسم الموظف", "التوقيع", "التاريخ"):
        p = doc.add_paragraph()
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(f"{label}: {'_' * 28}")
        set_run_font(run, size=11, color=GOLD_TEXT)

    add_heading(doc, "اعتماد الإدارة", size=12, color=GOLD_TEXT, space_before=10)
    for label in ("اسم المسؤول", "الوظيفة: المدير العام", "التوقيع", "ختم الشركة"):
        p = doc.add_paragraph()
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(f"{label}: {'_' * 24}")
        set_run_font(run, size=10.5, color=MUTED)


def build_employee_file(employee: Employee, role: RolePolicy, dest: Path) -> None:
    shutil.copy2(template_path(), dest)
    doc = Document(str(dest))

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(6)
    spacer.paragraph_format.space_after = Pt(4)

    add_heading(doc, "المهام والسياسات والالتزامات الوظيفية", size=17)
    add_meta_line(doc, "الشركة", ORG["name_ar_legal"])
    add_meta_line(doc, "المسمى الوظيفي", employee.job_title)
    add_meta_line(doc, "اسم الموظف", employee.name)
    if employee.nationality:
        add_meta_line(doc, "الجنسية", employee.nationality)
    add_meta_line(doc, "المرجع", f"{ORG['ref_prefix']}-HR-{employee.role_number:02d}")

    add_section(doc, "أولاً: المهام", role.tasks)
    add_section(doc, "ثانياً: السياسات والمسؤوليات", role.policies)
    add_section(doc, "ثالثاً: الالتزامات والواجبات", role.commitments)
    add_acknowledgment(doc, employee)

    doc.save(str(dest))


def main() -> None:
    lines = load_policy_lines()
    roles = parse_role_policies(lines)
    if not roles:
        print("Error: could not parse role policies.", file=sys.stderr)
        sys.exit(1)

    out = output_dir()
    created: list[Path] = []
    for employee in EMPLOYEES:
        role = roles.get(employee.role_number)
        if not role:
            print(f"Warning: no policy block for role #{employee.role_number}", file=sys.stderr)
            continue
        dest = out / safe_filename(employee.name, employee.job_title)
        build_employee_file(employee, role, dest)
        created.append(dest)
        print(f"Created: {dest}")

    print(f"\nDone — {len(created)} employee files in:\n{out}")


if __name__ == "__main__":
    main()
